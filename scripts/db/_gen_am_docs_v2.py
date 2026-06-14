"""
Per-AM outreach docs (Monday deliverable) per docs/design/PER_AM_DOC_SPEC.md.
Two sections: (1) Opportunities = this AM's deck cards (post-13.7 clean caps + 13.3 industry filter),
(2) Your major accounts = this AM's slice of the filter-aligned big-accounts top-100.
Verbatim opening note at top. Plain language, no jargon, no long dashes. NIC + LINDA only.
Inputs (read-only): card_am_assignment.json, outreach_cards_50.json, big_accounts.json.
Output: docs/data-insights/outreach_<am>.docx . Nothing sent.
"""
import os, io, sys, json, time
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
sys.path.insert(0, "backend")
from dotenv import load_dotenv
from supabase import create_client
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUTDIR = os.path.join(HERE, "..", "..", "docs", "data-insights")
os.makedirs(OUTDIR, exist_ok=True)
ASSIGN = json.load(open(os.path.join(HERE, "card_am_assignment.json"), encoding="utf-8"))
CARDS = {c["company_id"]: c for c in json.load(open(os.path.join(HERE, "outreach_cards_50.json"), encoding="utf-8"))["cards"]}
BIG = json.load(open(os.path.join(HERE, "big_accounts.json"), encoding="utf-8"))
GEN_DATE = "2026-06-14"
ONLY = ["Nic", "Linda"]
AM_FULL = {"Nic": "Nic Doyle", "Linda": "Linda D'Arcy"}   # to match big-accounts am_job

# trailing-12-month revenue per card company (the card's factors.revenue is ALL-TIME, not 12mo)
load_dotenv("backend/.env.production" if os.path.exists("backend/.env.production") else "backend/.env")
_sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])
_C8 = "241d7b99-f099-4557-96e5-212c4af10812"
_cids = list(CARDS.keys())
_sb.rpc("exec_sql", {"query": """CREATE OR REPLACE FUNCTION _tmp_cardrev(p uuid, ids uuid[]) RETURNS jsonb
LANGUAGE sql STABLE SET search_path=public AS $f$
SELECT coalesce(jsonb_object_agg(matched_company_id, r),'{}') FROM (
  SELECT matched_company_id, round(sum(cost_plus_price)::numeric,2) r FROM qb_operations
  WHERE client_id=p AND matched_company_id = ANY(ids) AND date_accepted >= DATE '2025-06-14'
  GROUP BY 1) t $f$;"""}).execute()
_sb.rpc("exec_sql", {"query": "NOTIFY pgrst,'reload schema'"}).execute(); time.sleep(2)
REV12 = {k: float(v) for k, v in (_sb.rpc("_tmp_cardrev", {"p": _C8, "ids": _cids}).execute().data or {}).items()}
_sb.rpc("exec_sql", {"query": "DROP FUNCTION IF EXISTS _tmp_cardrev(uuid,uuid[])"}).execute()

# Reorder-CADENCE metric for the rep-facing timing line. The deck measures gaps between individual
# order DATES, so a customer who places several jobs in the same fortnight shows a tiny median
# ("every 4 days"). Fix: collapse orders within BURST_DAYS into one ordering OCCASION (a project),
# then take the median/p90 gap BETWEEN occasions. Presentation-only (deck ranking unchanged).
BURST_DAYS = 14
_sb.rpc("exec_sql", {"query": f"""CREATE OR REPLACE FUNCTION _tmp_cadence(p uuid, ids uuid[]) RETURNS jsonb
LANGUAGE sql STABLE SET search_path=public AS $f$
WITH d AS (SELECT DISTINCT o.matched_company_id cid, o.date_accepted dt FROM qb_operations o
           WHERE o.client_id=p AND o.matched_company_id = ANY(ids) AND o.date_accepted IS NOT NULL),
mk AS (SELECT cid, dt, CASE WHEN lag(dt) OVER (PARTITION BY cid ORDER BY dt) IS NULL
                            OR dt - lag(dt) OVER (PARTITION BY cid ORDER BY dt) > {BURST_DAYS}
                       THEN 1 ELSE 0 END newocc FROM d),
occ AS (SELECT cid, dt, sum(newocc) OVER (PARTITION BY cid ORDER BY dt) oid FROM mk),
occd AS (SELECT cid, oid, min(dt) odt FROM occ GROUP BY cid, oid),
g AS (SELECT cid, odt - lag(odt) OVER (PARTITION BY cid ORDER BY odt) gap FROM occd),
agg AS (SELECT cid, (SELECT count(*) FROM occd o2 WHERE o2.cid=g.cid) n_occ,
               percentile_cont(0.5) WITHIN GROUP (ORDER BY gap) p50,
               percentile_cont(0.9) WITHIN GROUP (ORDER BY gap) p90
        FROM g GROUP BY cid)
SELECT coalesce(jsonb_object_agg(cid, jsonb_build_object('p50',p50,'p90',p90,'n_occ',n_occ)),'{{}}') FROM agg $f$;"""}).execute()
_sb.rpc("exec_sql", {"query": "NOTIFY pgrst,'reload schema'"}).execute(); time.sleep(2)
CADENCE = _sb.rpc("_tmp_cadence", {"p": _C8, "ids": _cids}).execute().data or {}
_sb.rpc("exec_sql", {"query": "DROP FUNCTION IF EXISTS _tmp_cadence(uuid,uuid[])"}).execute()

BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREY = RGBColor(0x70, 0x70, 0x70)
GREEN = RGBColor(0x1B, 0x7F, 0x3B)
AMBER = RGBColor(0xB5, 0x6A, 0x00)


def nd(s):
    return s.replace("—", "-").replace("–", "-").replace("―", "-") if isinstance(s, str) else s


def sanitize_doc(doc):
    def fix(paras):
        for p in paras:
            for r in p.runs:
                if r.text and any(ch in r.text for ch in ("—", "–", "―")):
                    r.text = nd(r.text)
    fix(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                fix(cell.paragraphs)


def shade(cell, hexfill):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_col_widths(table, widths):
    table.autofit = False; table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def add_runs(p, parts):
    for text, bold, italic, size, color in parts:
        r = p.add_run(nd(text)); r.bold = bold; r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color


def build_opening_note(doc, am, n_cards):
    h = doc.add_heading(f"Next-Best-Outreach for {am}", level=0)
    sub = doc.add_paragraph()
    add_runs(sub, [(f"Carbon8 - {n_cards} opportunity cards - generated {GEN_DATE} - "
                    "PROVISIONAL (final once mailbox data is complete)", False, True, 9.5, GREY)])

    box = doc.add_paragraph()
    add_runs(box, [("What changed since last time. ", True, False, None, BLUE),
                   ("The previous cards over-counted some categories because of how product tags were "
                    "stored: cello and laminate finishing was being counted as \"embellishment,\" which "
                    "inflated that category. That is now fixed at the platform level, so every pitch below "
                    "reads from corrected data.", False, False, None, None)])
    box2 = doc.add_paragraph()
    add_runs(box2, [("You will see a lot of Hard Cover / casebound suggestions, and that is the finding, "
                     "not a glitch. ", True, False, None, BLUE),
                    ("Once the tagging was clean, hard cover turned out to be the single biggest genuine gap "
                     "across your book-and-finishing customers: they already buy embellishment, soft cover and "
                     "finishing, but have not been sold casebound work, and customers with that exact profile "
                     "take it up at 3 to 4 times the base rate. Each card tells you whether the pitch ",
                     False, False, None, None),
                    ("fits the customer's industry", True, False, None, None),
                    (" (for example a property or trade-print client genuinely does books) or is ",
                     False, False, None, None),
                    ("purely statistical", True, False, None, None),
                    (" (worth a look, no industry signal either way). Where a strong-looking pitch did not fit "
                     "the customer's business, it is shown as ", False, False, None, None),
                    ("\"Not pitched\"", True, False, None, None),
                    (" with the reason, so you can see what the system set aside and why.", False, False, None, None)])


def fmt_money(v):
    return f"${round(v):,}"


DORMANT_DAYS = 365


def plain_timing(card):
    """Rep-facing timing in plain words. The 'typical reorder gap' uses the OCCASION-based cadence
    (same-project bursts collapsed) so it reads like a real reorder rhythm, not 'every 4 days'.
    days-since-last-order stays the actual last order; status is re-derived from the corrected gaps.
    Deck JSON unchanged."""
    te = card["back"]["timing_evidence"]
    days = te.get("days_since_last_order")
    seas = te.get("seasonal") or {}
    cad = CADENCE.get(card["company_id"]) or {}
    p50 = cad.get("p50")
    p90 = cad.get("p90")
    n_occ = cad.get("n_occ") or 0
    typ = f"they typically reorder about every {round(p50)} days" if p50 is not None else None

    if days is None:
        base = "No order-date history to read a reorder timing from."
    elif days > DORMANT_DAYS:
        base = (f"Dormant: last ordered {days} days ago (about {days/365:.1f} years), so this is a "
                f"reactivation, not a warm reorder.")
    elif n_occ < 2 or p50 is None:
        base = f"Only one ordering occasion on record ({days} days ago), so there's no reorder rhythm to read yet."
    elif p90 is not None and days >= p90:
        base = f"Overdue: last ordered {days} days ago; {typ}, so they're past due."
    elif days >= p50:
        base = f"Due soon: last ordered {days} days ago; {typ}, so they're about due."
    else:
        base = f"In cadence: last ordered {days} days ago; {typ}, so they're not due yet."

    if seas.get("approaching") and seas.get("peak_month"):
        base += f" They also tend to pick up around {seas['peak_month']}, which is coming up."
    return base


def pattern_line(card):
    gr = card["front"]["context"]["gap_rationale"]
    if not gr.get("pitched"):
        return None
    X = gr["predicted_by"]; Y = gr["pitched"]
    return (f"{gr['confidence_pct']:.0f}% of customers who buy {X} also buy {Y}, about {gr['lift']:.1f} times more "
            f"likely than the average customer ({gr['base_rate_pct']:.0f}% normally). {gr['fact'].capitalize()}.")


def build_opportunities(doc, am, recs):
    doc.add_heading("Section 1 - Opportunities", level=1)
    doc.add_paragraph("One card per customer. The numbers come straight from QuickBase order history; "
                      "nothing is rounded up. Mark each one in the feedback line.", style="Intense Quote")
    for idx, r in enumerate(recs, 1):
        card = CARDS[r["company_id"]]
        f = card["front"]; inf = f["industry_fit"]
        rev12 = REV12.get(r["company_id"], 0)
        rev_txt = (f"{fmt_money(rev12)} in the last 12 months" if rev12 > 0 else "no orders in the last 12 months")
        ind = inf["industry"]
        h = doc.add_heading("", level=3)
        add_runs(h, [(f"{idx}. {f['customer']}  ", False, False, None, None),
                     (f"[{ind}]", True, False, 10, GREY),
                     (f"   {rev_txt}   (deck #{card['rank']} of 50)", False, True, 9, GREY)])

        # contact
        p = doc.add_paragraph(); add_runs(p, [("Who to contact: ", True, False, None, None), (f["to_whom"], False, False, None, None)])
        cn = card["back"]["contact_evidence"].get("contrast_note")
        if cn and "NOT the loudest" in cn:
            pc = doc.add_paragraph(); add_runs(pc, [("   Note: ", True, False, 9, AMBER), (cn, False, True, 9, GREY)])

        # pitch
        pp = doc.add_paragraph()
        if inf["kept_pitch"]:
            add_runs(pp, [("Pitch: ", True, False, None, None), (inf["kept_pitch"], False, False, None, GREEN)])
        else:
            add_runs(pp, [("Pitch: ", True, False, None, None), (f["pitch"], False, False, None, GREY)])

        pat = pattern_line(card)
        if pat:
            pp2 = doc.add_paragraph(); add_runs(pp2, [("Pattern: ", True, False, None, None), (pat, False, False, None, None)])

        # fits their business (state neutral plainly; the deck's fit line already carries its own label)
        pf = doc.add_paragraph()
        fitv = inf["fit"]
        if fitv.startswith("Fits their business: "):
            add_runs(pf, [("Fits their business: ", True, False, None, GREEN), (fitv[len("Fits their business: "):], False, False, None, None)])
        else:
            add_runs(pf, [("Industry fit: ", True, False, None, None), (fitv, False, False, None, None)])

        # not pitched (suppression) - prominent where present
        if inf.get("not_pitched"):
            pn = doc.add_paragraph()
            add_runs(pn, [("Not pitched: ", True, False, None, AMBER), (inf["not_pitched"], False, False, None, None)])

        pt = doc.add_paragraph(); add_runs(pt, [("Timing: ", True, False, None, None), (plain_timing(card), False, False, None, None)])

        fb = doc.add_paragraph()
        add_runs(fb, [("Your call:   ", True, False, None, BLUE),
                      ("[ ] Makes sense      [ ] They already buy this      [ ] Wrong      ", False, False, None, None),
                      ("notes: ", False, True, 9, GREY), ("_" * 30, False, False, None, RGBColor(0xA0, 0xA0, 0xA0))])


def build_major_accounts(doc, am, accounts, shifts, drops):
    doc.add_page_break()
    doc.add_heading("Section 2 - Your major accounts", level=1)
    n = len(accounts)
    sat = sum(1 for a in accounts if a["status"] in ("saturated", "retention-focus"))
    gap = [a for a in accounts if a["status"] == "has-real-gap"]
    intro = doc.add_paragraph()
    add_runs(intro, [(f"Your {n} biggest accounts by revenue over the last 12 months. ", False, False, None, None),
                     (f"Most are well covered ({sat} buy their full range, no real gap), which is a good thing: "
                      f"these are relationships to protect. ", False, False, None, None),
                     (f"{len(gap)} of them still have a genuine gap worth a conversation, listed first.",
                      False, False, None, None)])

    # surface any shift in this AM's slice as a positive
    for cust, ind, raw, kept in shifts:
        ps = doc.add_paragraph()
        add_runs(ps, [("Smart catch: ", True, False, None, GREEN),
                      (f"for {cust} ({ind}), the system swapped a generic {raw} suggestion for {kept}, which "
                       f"is what {ind} clients actually buy.", False, False, None, None)])
    for d in drops:
        pdp = doc.add_paragraph()
        add_runs(pdp, [("Smart catch: ", True, False, None, GREEN),
                       (f"we did NOT flag {d['suppressed_gap']} for {d['customer']} ({d['industry']}) even though "
                        f"the raw numbers suggested it, because {d['industry']} firms rarely buy it "
                        f"({d['industry_buy_rate_pct']}%).", False, False, None, None)])

    headers = ["Customer", "Industry", "12-mo revenue", "Status", "Note"]
    widths = [Inches(1.7), Inches(1.3), Inches(1.05), Inches(1.0), Inches(2.15)]
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        rr = hdr[i].paragraphs[0].add_run(htext); rr.bold = True; rr.font.size = Pt(8)
        shade(hdr[i], "D5E8F0"); hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    order = {"has-real-gap": 0, "gap-suppressed": 1, "saturated": 2, "retention-focus": 3}
    for a in sorted(accounts, key=lambda x: (order.get(x["status"], 9), -x["rev_12mo"])):
        if a["status"] == "has-real-gap":
            note = f"Real gap: {a['pitch'].split(' (')[0]}" + (f" ({a['industry_fit']})" if a.get("industry_fit") else "")
            label = "Real gap"
        elif a["status"] == "saturated":
            note = "Buys your full range, no real gap. Retention relationship, protect it."; label = "Saturated"
        elif a["status"] == "gap-suppressed":
            note = "Numbers suggested a gap that does not fit their industry. Treat as retention."; label = "No fit gap"
        else:
            note = "No clear new product to add. Reorder / keep-warm relationship."; label = "Retention"
        cells = table.add_row().cells
        for i, v in enumerate([a["customer"], a["industry"], fmt_money(a["rev_12mo"]), label, note]):
            cp = cells[i].paragraphs[0]; run = cp.add_run(nd(str(v))); run.font.size = Pt(8)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_col_widths(table, widths)


def build_doc(am, recs):
    recs = sorted(recs, key=lambda r: CARDS[r["company_id"]]["rank"])
    full = AM_FULL.get(am, am)
    accounts = [a for a in BIG["accounts"] if (a.get("am") or "").startswith(am) or a.get("am") == full]
    shifts = [(s["customer"], s["industry"], s["raw_gap"], s["fitting_gap"]) for s in BIG.get("shifted_gaps", [])
              if any(a["customer"] == s["customer"] for a in accounts)]
    drops = [d for d in BIG.get("dropped_implausible_gaps", []) if (d.get("am") or "").startswith(am) or d.get("am") == full]

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.5))
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    build_opening_note(doc, am, len(recs))
    build_opportunities(doc, am, recs)
    build_major_accounts(doc, am, accounts, shifts, drops)
    sanitize_doc(doc)
    fn = os.path.join(OUTDIR, f"outreach_{am.lower()}.docx")
    doc.save(fn)
    return fn, len(recs), len(accounts), len(shifts), len(drops)


if __name__ == "__main__":
    per_am = {}
    for c in ASSIGN["cards"]:
        if c.get("am"):
            per_am.setdefault(c["am"], []).append(c)
    print("Generating NIC + LINDA only (per spec sequence)...\n")
    for am in ONLY:
        recs = per_am.get(am, [])
        if not recs:
            print(f"  {am}: no cards"); continue
        try:
            fn, n, na, ns, ndp = build_doc(am, recs)
            # pitch-type mix
            mix = {}
            for r in recs:
                kp = CARDS[r["company_id"]]["front"]["industry_fit"]["kept_pitch"] or "retention/none"
                mix[kp] = mix.get(kp, 0) + 1
            supp = sum(1 for r in recs if CARDS[r["company_id"]]["front"]["industry_fit"].get("not_pitched"))
            print(f"  {am}: {n} cards -> {os.path.basename(fn)}")
            print(f"     pitch mix: {dict(sorted(mix.items(), key=lambda kv:-kv[1]))}")
            print(f"     major-accounts slice: {na} | shifts surfaced: {ns} | suppressions surfaced: {ndp} | "
                  f"cards with a 'Not pitched' line: {supp}")
        except PermissionError:
            print(f"  SKIPPED outreach_{am.lower()}.docx - file is open in Word, close and re-run")
