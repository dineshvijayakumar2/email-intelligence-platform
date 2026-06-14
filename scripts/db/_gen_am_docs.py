"""
Step 3 — one Word doc per AM (>=1 assigned customer). Read-only inputs:
  card_am_assignment.json (AM + grounding) + outreach_cards_50.json (full card detail).
Identical top section (what this is / how to read / tiers / honest notes), a per-AM summary
table with a blank AM-feedback column, and supporting detail per customer with a blank
feedback prompt. Honest text where contact/timing is missing (no mailbox). -> outreach_<am>.docx
"""
import os, io, sys, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUTDIR = os.path.join(HERE, "..", "..", "docs", "data-insights")
os.makedirs(OUTDIR, exist_ok=True)
ASSIGN = json.load(open(os.path.join(HERE, "card_am_assignment.json"), encoding="utf-8"))
CARDS = {c["company_id"]: c for c in json.load(open(os.path.join(HERE, "outreach_cards_50.json"), encoding="utf-8"))["cards"]}
GEN_DATE = "2026-06-12"


def nd(s):
    """Strip long dashes (em / en / horizontal bar) -> hyphen; keep normal hyphens."""
    if not isinstance(s, str):
        return s
    return s.replace("—", "-").replace("–", "-").replace("―", "-")


def sanitize_doc(doc):
    """Post-pass: replace any long dash in every run (covers JSON-sourced text too)."""
    def fix(paras):
        for p in paras:
            for r in p.runs:
                if r.text and any(ch in r.text for ch in ("—", "–", "―")):
                    r.text = nd(r.text)
    fix(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                fix(cell.paragraphs)

TIER_COLOR = {"STRONG": RGBColor(0x1B, 0x7F, 0x3B), "MODERATE": RGBColor(0xB5, 0x6A, 0x00),
              "WEAK": RGBColor(0x70, 0x70, 0x70)}


def tier_of(conf, lift):
    if conf is None or lift is None:
        return "WEAK"
    if conf >= 0.50 and lift >= 1.5:
        return "STRONG"
    if conf >= 0.35 and lift >= 1.3:
        return "MODERATE"
    return "WEAK"


def shade(cell, hexfill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def add_runs(p, parts):
    """parts = list of (text, bold, italic, size_pt, color)"""
    for text, bold, italic, size, color in parts:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color


def build_top(doc, am, n):
    h = doc.add_heading(f"Next-Best-Outreach — {am}", level=0)
    sub = doc.add_paragraph()
    add_runs(sub, [(f"Carbon8 · {n} customer{'s' if n != 1 else ''} · generated {GEN_DATE} · "
                    "PROVISIONAL (final once mailbox data is complete)", False, True, 9.5, RGBColor(0x60, 0x60, 0x60))])

    doc.add_heading("What this is", level=2)
    doc.add_paragraph(
        "Data-derived outreach recommendations — one card per customer: what to pitch, to whom, "
        "with what context, and at what timing. These are read-only suggestions built from QuickBase "
        "order history (the ground truth) and email engagement. They are a starting point for your "
        "judgement, not a directive — please mark up the feedback column and prompts.")

    doc.add_heading("How to read Confidence & Lift", level=2)
    for t in [
        ("Confidence", " — of customers who already buy the predictor product, the percentage who "
         "also buy the pitched product. Higher = more of their peers do it. Stated as the real number, never rounded up."),
        ("Lift", " — how much more likely this customer is to want the product than the average customer. "
         "Lift 1.0 = no signal (almost everyone buys it); lift 1.5+ = a strong, genuine adjacency."),
        ("Watch-out", " — a high-confidence / low-lift pitch (e.g. 90% but lift ~1.0) just means “almost "
         "everyone buys it” — it's weak as a targeted pitch, which is why those land in the WEAK tier below."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, [(t[0], True, False, None, None), (t[1], False, False, None, None)])

    doc.add_heading("Tiers (lift-aware)", level=2)
    for label, desc in [
        ("STRONG", "confidence ≥50% AND lift ≥1.5 — high confidence and a genuine adjacency. Lead with these."),
        ("MODERATE", "confidence ≥35% AND lift ≥1.3 — worth a look, softer signal."),
        ("WEAK", "low confidence (<35%), or near-universal/low-lift, or a fallback pitch — context only, use your judgement."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, [(label, True, False, None, TIER_COLOR[label]), (f" — {desc}", False, False, None, None)])

    doc.add_heading("Honest notes", level=2)
    for note in [
        "QuickBase orders are the ground truth. Every pitch was verified as a genuine gap — the customer "
        "has zero orders in the pitched product — including an embellishment finish-tag check (so we don't "
        "pitch “embellishment” to someone who already foils or embosses within their jobs).",
        "Where there is no AM mailbox or no email captured for a customer, the contact and/or timing may be "
        "missing. The card says so explicitly — please don't read a blank as “no relationship,” it means "
        "“not in our email data.” Confirm the contact before reaching out.",
        "Confidence is the actual rule percentage; timing is read from the customer's own order rhythm "
        "(how overdue they are vs their own typical gap between orders).",
    ]:
        doc.add_paragraph(note, style="List Bullet")


def fmt_conf(conf):
    return f"{conf*100:.0f}%" if conf is not None else "—"


def build_summary_table(doc, am, rows):
    doc.add_heading(f"Summary — {am}'s {len(rows)} customer{'s' if len(rows) != 1 else ''}", level=2)
    headers = ["#", "Customer", "Pitch", "Confidence", "Lift", "Timing", "Tier", "AM feedback"]
    widths = [Inches(0.3), Inches(1.4), Inches(1.0), Inches(0.7), Inches(0.45),
              Inches(1.05), Inches(0.6), Inches(1.7)]   # portrait A4, ~7.2in content
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        r0 = hdr[i].paragraphs[0].add_run(htext)
        r0.bold = True
        r0.font.size = Pt(8)
        shade(hdr[i], "D5E8F0")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for idx, r in enumerate(rows, 1):
        card = r["_card"]
        cells = table.add_row().cells
        vals = [str(idx), r["customer"], r["_pitch_short"], fmt_conf(r["_conf_frac"]),
                (f"{r['lift']:.2f}" if r["lift"] else "-"),
                r["_timing_short"], r["_tier"], ""]
        for i, v in enumerate(vals):
            cp = cells[i].paragraphs[0]
            run = cp.add_run(v)
            run.font.size = Pt(8)
            if i == 6:  # tier colour
                run.bold = True
                run.font.color.rgb = TIER_COLOR.get(v, RGBColor(0, 0, 0))
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_col_widths(table, widths)


def contact_line(card, grounding):
    vc = card["back"]["contact_evidence"]["value_contact"]
    cs = grounding["contact_status"]
    if cs == "engaged":
        nm = vc["name"] or vc["email"]
        fresh = "stale >90d — confirm contact" if vc["stale"] else "active"
        extra = f"last active {vc['last_active']}" if vc["last_active"] else "no recent activity"
        return (f"{nm} <{vc['email']}> — {fresh} ({extra}). Won ${vc['won_val']:,} over "
                f"{vc['won']}/{vc['quotes']} quotes (strike {vc['strike']}%), email {vc['em_from']} in / {vc['em_to']} out, "
                f"{'two-way' if vc['two_way'] else 'mostly one-way'}.")
    if cs == "qb_only_no_email":
        nm = vc["name"] or vc["email"]
        return (f"{nm} <{vc['email']}> — identified from QuickBase quotes (won ${vc['won_val']:,}), but NO email "
                "activity is captured for this account (no AM mailbox / not synced). Confirm this is the right contact before outreach.")
    return "No QuickBase-won value contact identified — confirm the decision-maker before outreach."


def timing_line(card, grounding):
    te = card["back"]["timing_evidence"]
    if not grounding["timing_available"]:
        return "Timing unavailable — no order-date history on record."
    bits = card["front"]["timing"]
    extra = f"  (p50 gap {te['p50_gap_days']}d, p90 gap {te['p90_gap_days']}d, last order {te['days_since_last_order']}d ago, {te['n_orders']} orders)"
    seas = te.get("seasonal") or {}
    if seas.get("approaching"):
        extra += f"  · approaching {seas.get('peak_month')} seasonal peak"
    return bits + extra


def build_detail(doc, rows):
    doc.add_heading("Detail & your feedback", level=2)
    for idx, r in enumerate(rows, 1):
        card = r["_card"]
        g = r["grounding"]
        h = doc.add_heading("", level=3)
        add_runs(h, [(f"{idx}. {r['customer']}  ", False, False, None, None),
                     (f"[{r['_tier']}]", True, False, 11, TIER_COLOR[r["_tier"]]),
                     (f"   (overall #{r['rank']} of 50)", False, True, 9, RGBColor(0x80, 0x80, 0x80))])

        pe = card["back"]["pitch_evidence"]
        p = doc.add_paragraph()
        add_runs(p, [("Pitch: ", True, False, None, None), (card["front"]["pitch"], False, False, None, None)])
        p2 = doc.add_paragraph()
        add_runs(p2, [("Why: ", True, False, None, None),
                      (card["front"]["confidence_stated"], False, False, None, None)])
        if pe.get("gap_confirmation"):
            gc = pe["gap_confirmation"]
            p3 = doc.add_paragraph()
            add_runs(p3, [("Gap check: ", True, False, None, None),
                          (f"{gc['company_jobs_in_pitch']} existing orders in this product — "
                           f"{'genuine gap confirmed' if gc['is_genuine_zero_order_gap'] else 'NOT a clean gap'}.",
                           False, False, None, None)])

        p4 = doc.add_paragraph()
        add_runs(p4, [("To whom: ", True, False, None, None), (contact_line(card, g), False, False, None, None)])

        lo = card["front"]["context"]["last_order"]
        ctx = (f"last order {lo['capability']} on {lo['date']}; " if lo else "")
        ctx += "basket — " + ", ".join(card["front"]["context"]["basket_summary"])
        p5 = doc.add_paragraph()
        add_runs(p5, [("Context: ", True, False, None, None), (ctx, False, False, None, None)])

        p6 = doc.add_paragraph()
        add_runs(p6, [("Timing: ", True, False, None, None), (timing_line(card, g), False, False, None, None)])

        fb = doc.add_paragraph()
        add_runs(fb, [("Your assessment: ", True, False, None, RGBColor(0x1F, 0x49, 0x7D)),
                      ("_" * 70, False, False, None, RGBColor(0xA0, 0xA0, 0xA0))])
        cue = doc.add_paragraph()
        add_runs(cue, [("Does the recommendation make sense?  Do they already buy it?  Is the contact right?  "
                        "Better timing or a better product to lead with?", False, True, 9, RGBColor(0x80, 0x80, 0x80))])


def build_doc(am, recs, slug=None):
    recs = sorted(recs, key=lambda r: r["rank"])
    for r in recs:
        card = CARDS[r["company_id"]]
        r["_card"] = card
        r["_conf_frac"] = card["factors"]["confidence"]   # FRACTION (front.confidence is already a percent)
        r["_tier"] = tier_of(r["_conf_frac"], r["lift"])
        r["_pitch_short"] = (card["front"]["pitch"] or "—").split(" —")[0].strip()
        if "no high-confidence" in (card["front"]["pitch"] or ""):
            r["_pitch_short"] = "retention/reorder"
        r["_timing_short"] = card["back"]["timing_evidence"]["status"].replace("_", " ")
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)   # A4 portrait
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.5))                                # tight margins
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    build_top(doc, am, len(recs))
    build_summary_table(doc, am, recs)
    doc.add_page_break()
    build_detail(doc, recs)
    sanitize_doc(doc)                                              # no long dashes
    fn = os.path.join(OUTDIR, f"outreach_{(slug or am.lower())}.docx")
    doc.save(fn)
    return fn, len(recs)


if __name__ == "__main__":
    # group assigned cards by AM (only the six; skip flagged/unassigned)
    per_am = {}
    for c in ASSIGN["cards"]:
        if c.get("am"):
            per_am.setdefault(c["am"], []).append(c)

    made = []
    for am in ["Nic", "Linda", "Ehab", "Kenneth", "Mary", "Peter"]:
        if am in per_am and per_am[am]:
            try:
                fn, n = build_doc(am, per_am[am])
                made.append((am, os.path.basename(fn), n))
                print(f"  wrote {os.path.basename(fn)}  ({n} customers)")
            except PermissionError:
                print(f"  SKIPPED outreach_{am.lower()}.docx - file is open/locked (close it in Word, then re-run)")

    print("\nDocs generated:", len(made))
    for am, fn, n in made:
        print(f"  {am:<9} {fn:<22} {n} customers")
